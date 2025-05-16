import streamlit as st
from docx import Document
from ebooklib import epub
from fpdf import FPDF
from io import BytesIO
import base64

st.set_page_config(page_title="AI Book Formatter", layout="wide")
st.title("📖 AI Book Formatter for Amazon Publishing")

# Sidebar for input
st.sidebar.header("Book Settings")
title = st.sidebar.text_input("Book Title")
author = st.sidebar.text_input("Author")
subtitle = st.sidebar.text_input("Subtitle (optional)")
trim_size = st.sidebar.selectbox("Trim Size", ["6x9", "5.5x8.5", "8.5x11"])
formats = st.sidebar.multiselect("Output Formats", ["Print (.docx)", "Print (.pdf)", "Kindle (.epub)"])
include_toc = st.sidebar.checkbox("Include Table of Contents", value=True)
include_front_matter = st.sidebar.checkbox("Add Front & Back Matter", value=True)
clean_grammar = st.sidebar.checkbox("AI Clean & Format (Simulated)", value=True)

# File uploader
uploaded_file = st.file_uploader("Upload Your Book Content (.docx or .txt)", type=["docx", "txt"])

if uploaded_file and title and author:
    raw_text = ""
    if uploaded_file.name.endswith(".txt"):
        raw_text = uploaded_file.read().decode("utf-8")
    else:
        doc = Document(uploaded_file)
        raw_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])

    st.success("File uploaded and parsed.")

    if clean_grammar:
        raw_text = raw_text.replace("\n\n", "\n")

    chapters = raw_text.split("Chapter ")[1:] if "Chapter " in raw_text else [raw_text]
    chapters = ["Chapter " + c.strip() for c in chapters]

    def generate_docx():
        output = BytesIO()
        docx = Document()
        if include_front_matter:
            docx.add_heading(title, 0)
            if subtitle:
                docx.add_paragraph(subtitle)
            docx.add_paragraph(f"by {author}")
            docx.add_page_break()
        if include_toc:
            docx.add_heading("Table of Contents", level=1)
            for i, ch in enumerate(chapters):
                docx.add_paragraph(f"Chapter {i+1}", style='List Number')
            docx.add_page_break()
        for ch in chapters:
            docx.add_heading(ch.split("\n")[0], level=1)
            for p in ch.split("\n")[1:]:
                docx.add_paragraph(p)
        docx.save(output)
        return output.getvalue()

    def generate_epub():
        book = epub.EpubBook()
        book.set_identifier("id123456")
        book.set_title(title)
        book.set_language("en")
        book.add_author(author)
        epub_chapters = []
        for i, ch in enumerate(chapters):
            chap = epub.EpubHtml(title=f"Chapter {i+1}", file_name=f'chap_{i+1}.xhtml', lang='en')
            content = f'<h1>Chapter {i+1}</h1><p>{ch}</p>'
            chap.content = content
            book.add_item(chap)
            epub_chapters.append(chap)
        book.toc = tuple(epub_chapters)
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        style = 'BODY { font-family: Times, serif; }'
        nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)
        book.add_item(nav_css)
        epub.write_epub("output.epub", book)
        with open("output.epub", "rb") as f:
            return f.read()

    def generate_pdf():
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Times", 'B', 16)
        pdf.multi_cell(0, 10, title.encode('latin-1', 'replace').decode('latin-1'))
        if subtitle:
            pdf.set_font("Times", '', 12)
            pdf.multi_cell(0, 10, subtitle.encode('latin-1', 'replace').decode('latin-1'))
        pdf.set_font("Times", '', 10)
        pdf.multi_cell(0, 10, f"by {author}".encode('latin-1', 'replace').decode('latin-1'))
        pdf.ln(10)
        for ch in chapters:
            pdf.add_page()
            pdf.set_font("Times", 'B', 14)
            pdf.multi_cell(0, 10, ch.split("\n")[0].encode('latin-1', 'replace').decode('latin-1'))
            pdf.set_font("Times", '', 12)
            for p in ch.split("\n")[1:]:
                pdf.multi_cell(0, 10, p.encode('latin-1', 'replace').decode('latin-1'))
        pdf_output = BytesIO()
        pdf.output(pdf_output)
        return pdf_output.getvalue()

    if "Print (.docx)" in formats:
        docx_bytes = generate_docx()
        st.download_button("📄 Download DOCX", docx_bytes, file_name="formatted_book.docx")

    if "Kindle (.epub)" in formats:
        epub_bytes = generate_epub()
        st.download_button("📱 Download EPUB", epub_bytes, file_name="formatted_book.epub")

    if "Print (.pdf)" in formats:
        pdf_bytes = generate_pdf()
        st.download_button("🖨️ Download PDF", pdf_bytes, file_name="formatted_book.pdf")

else:
    st.warning("Please upload a file and fill in all required fields.")
